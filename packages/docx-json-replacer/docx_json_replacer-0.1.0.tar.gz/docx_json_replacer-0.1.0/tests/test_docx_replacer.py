import pytest
import json
import os
from docx import Document
from docx_json_replacer.docx_replacer import DocxReplacer, replace_docx_template


class TestDocxReplacer:
    def test_init(self, sample_docx):
        replacer = DocxReplacer(sample_docx)
        assert replacer.docx_path == sample_docx
        assert replacer.document is not None

    def test_replace_from_json(self, docx_replacer, test_data, output_docx):
        docx_replacer.replace_from_json(test_data)
        docx_replacer.save(output_docx)
        
        # Verify the replacements were made
        result_doc = Document(output_docx)
        
        # Check paragraphs
        paragraphs_text = [p.text for p in result_doc.paragraphs]
        assert "Hello John Doe!" in paragraphs_text
        assert "Date: 2025-06-25" in paragraphs_text
        assert "Company: Example Corp" in paragraphs_text
        assert "Position: Software Engineer" in paragraphs_text
        
        # Check tables
        table = result_doc.tables[0]
        assert table.cell(0, 1).text == "John Doe"
        assert table.cell(1, 1).text == "Software Engineer"

    def test_replace_from_json_file(self, docx_replacer, sample_json_file, output_docx):
        docx_replacer.replace_from_json_file(sample_json_file)
        docx_replacer.save(output_docx)
        
        # Verify the replacements were made
        result_doc = Document(output_docx)
        paragraphs_text = [p.text for p in result_doc.paragraphs]
        assert "Hello John Doe!" in paragraphs_text
        assert "Date: 2025-06-25" in paragraphs_text

    def test_save(self, docx_replacer, test_data, output_docx):
        docx_replacer.replace_from_json(test_data)
        docx_replacer.save(output_docx)
        
        assert os.path.exists(output_docx)
        
        # Verify the saved document is valid
        saved_doc = Document(output_docx)
        assert len(saved_doc.paragraphs) > 0

    def test_placeholder_not_found(self, docx_replacer, output_docx):
        data = {"nonexistent.key": "value"}
        docx_replacer.replace_from_json(data)
        docx_replacer.save(output_docx)
        
        # Document should still be valid even if no replacements were made
        result_doc = Document(output_docx)
        assert len(result_doc.paragraphs) > 0

    def test_empty_json_data(self, docx_replacer, output_docx):
        docx_replacer.replace_from_json({})
        docx_replacer.save(output_docx)
        
        # Document should still be valid
        result_doc = Document(output_docx)
        assert len(result_doc.paragraphs) > 0

    def test_partial_replacement(self, docx_replacer, output_docx):
        # Only replace some placeholders
        partial_data = {
            "input.name": "Jane Smith",
            "input.company": "New Corp"
        }
        docx_replacer.replace_from_json(partial_data)
        docx_replacer.save(output_docx)
        
        result_doc = Document(output_docx)
        paragraphs_text = [p.text for p in result_doc.paragraphs]
        assert "Hello Jane Smith!" in paragraphs_text
        assert "Company: New Corp" in paragraphs_text
        # Unreplaced placeholders should remain
        assert "Date: {{input.date}}" in paragraphs_text
        assert "Position: {{input.position}}" in paragraphs_text


class TestReplaceDocxTemplate:
    def test_replace_docx_template_function(self, sample_docx, test_data, output_docx):
        replace_docx_template(sample_docx, test_data, output_docx)
        
        assert os.path.exists(output_docx)
        
        # Verify the replacements were made
        result_doc = Document(output_docx)
        paragraphs_text = [p.text for p in result_doc.paragraphs]
        assert "Hello John Doe!" in paragraphs_text
        assert "Date: 2025-06-25" in paragraphs_text
        assert "Company: Example Corp" in paragraphs_text
        assert "Position: Software Engineer" in paragraphs_text


class TestEdgeCases:
    def test_multiple_same_placeholder_in_paragraph(self, sample_docx, output_docx):
        # Create a document with multiple instances of the same placeholder
        doc = Document()
        doc.add_paragraph("{{input.name}} works at {{input.company}}. {{input.name}} is a {{input.position}}.")
        doc.save(sample_docx)
        
        replacer = DocxReplacer(sample_docx)
        test_data = {
            "input.name": "Alice",
            "input.company": "Tech Corp",
            "input.position": "Developer"
        }
        replacer.replace_from_json(test_data)
        replacer.save(output_docx)
        
        result_doc = Document(output_docx)
        paragraph_text = result_doc.paragraphs[0].text
        assert paragraph_text == "Alice works at Tech Corp. Alice is a Developer."

    def test_numeric_values(self, docx_replacer, output_docx):
        numeric_data = {
            "input.age": 30,
            "input.salary": 75000.50,
            "input.years": 5
        }
        
        # Add numeric placeholders to the document
        docx_replacer.document.add_paragraph("Age: {{input.age}}")
        docx_replacer.document.add_paragraph("Salary: {{input.salary}}")
        docx_replacer.document.add_paragraph("Experience: {{input.years}} years")
        
        docx_replacer.replace_from_json(numeric_data)
        docx_replacer.save(output_docx)
        
        result_doc = Document(output_docx)
        paragraphs_text = [p.text for p in result_doc.paragraphs]
        assert "Age: 30" in paragraphs_text
        assert "Salary: 75000.5" in paragraphs_text
        assert "Experience: 5 years" in paragraphs_text