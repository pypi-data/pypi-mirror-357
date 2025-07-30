import json
from typing import Dict, Any
from docx import Document


class DocxReplacer:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.document = Document(docx_path)
    
    def replace_from_json(self, json_data: Dict[str, Any]) -> None:
        for paragraph in self.document.paragraphs:
            for key, value in json_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in json_data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
    
    def replace_from_json_file(self, json_path: str) -> None:
        with open(json_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        self.replace_from_json(json_data)
    
    def save(self, output_path: str) -> None:
        self.document.save(output_path)


def replace_docx_template(docx_path: str, json_data: Dict[str, Any], output_path: str) -> None:
    replacer = DocxReplacer(docx_path)
    replacer.replace_from_json(json_data)
    replacer.save(output_path)