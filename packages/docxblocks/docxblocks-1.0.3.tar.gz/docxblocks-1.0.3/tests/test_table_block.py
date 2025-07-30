import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_table_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["A", "B"],
                "rows": [["1", "2"], ["3", "4"]]
            },
            "style": {
                "header_styles": {"bold": True},
                "column_styles": {1: {"align": "center"}},
                "row_styles": {1: {"bg_color": "FFF8DC"}},
                "cell_styles": {(1, 1): {"font_color": "FF0000"}}
            }
        }
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    # Check that table exists and has correct number of rows
    tables = doc2.tables
    assert len(tables) == 1
    assert len(tables[0].rows) == 3  # header + 2 rows 