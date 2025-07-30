import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT

def test_image_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {"type": "image", "path": "nonexistent.png", "style": {"max_width": "2in"}}
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    found = any(DEFAULT_EMPTY_VALUE_TEXT in p.text for p in doc2.paragraphs)
    assert found, "Image placeholder text not found in output docx" 