"""
Table Builder Module

This module provides the TableBuilder class for rendering table blocks in Word documents.
It handles table creation with headers, rows, column styling, and cell formatting.
"""

from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxblocks.constants import DEFAULT_EMPTY_VALUE_TEXT, DEFAULT_EMPTY_VALUE_STYLE
from docxblocks.utils.styles import set_paragraph_alignment


class TableBuilder:
    """
    Builder class for rendering table blocks in Word documents.
    
    This builder handles table creation with headers, rows, and comprehensive
    styling options. It supports column widths, header styling, column styling,
    and individual cell styling. Empty content, headers, and cells are replaced
    with consistent placeholders.
    
    The builder creates tables with the "Table Grid" style by default and
    supports various styling options through the style_kwargs parameter.
    """
    
    @staticmethod
    def build(doc, placeholder=None, content=None, parent=None, index=None, **style_kwargs):
        """
        Build and render a table block in the document.
        
        This method processes table content, creates the table structure,
        applies styling, and handles empty values with placeholders.
        
        Args:
            doc: The python-docx Document object
            placeholder: Placeholder text (unused in current implementation)
            content: Dictionary containing table data with keys:
                - headers: List of header strings
                - rows: List of row data (list of lists)
            parent: The parent XML element where content will be inserted
            index: The insertion index within the parent element
            **style_kwargs: Additional styling options including:
                - column_widths: List of width fractions for columns
                - row_widths: List of height fractions for rows
                - header_styles: Dictionary of header styling
                - column_styles: Dictionary of column styling by index
                - row_styles: Dictionary of row styling by index
                - cell_styles: Dictionary of cell styling by (row, col) tuple
        """
        if parent is None or index is None:
            return

        # Ensure style_kwargs is always a dictionary
        style_kwargs = style_kwargs or {}

        # Handle empty content with placeholder
        if not content:
            para = doc.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            # Apply placeholder style
            run = para.runs[0]
            run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
            if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])
            parent.insert(index, para._element)
            return

        headers = content.get("headers", [])
        rows = content.get("rows", [])

        # Handle empty headers and rows
        if not headers and not rows:
            para = doc.add_paragraph(DEFAULT_EMPTY_VALUE_TEXT)
            # Apply placeholder style
            run = para.runs[0]
            run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
            if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])
            parent.insert(index, para._element)
            return

        num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
        table = doc.add_table(rows=0, cols=num_cols)
        table.style = "Table Grid"

        column_widths = style_kwargs.get("column_widths")
        row_widths = style_kwargs.get("row_widths")

        # Apply column widths
        if column_widths:
            total_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            for i, fraction in enumerate(column_widths):
                if i < len(table.columns):
                    table.columns[i].width = int(total_width * fraction)

        if headers:
            row = table.add_row().cells
            # Apply row height if specified
            if row_widths and len(row_widths) > 0:
                row_height = row_widths[0]
                if row_height:
                    table.rows[0].height = int(row_height * 914400)  # Convert to EMUs
            
            for i, header_text in enumerate(headers):
                cell = row[i]
                para = cell.paragraphs[0]
                # Handle empty header text
                header_display = header_text.strip() if header_text else DEFAULT_EMPTY_VALUE_TEXT
                run = para.add_run(str(header_display))
                run.font.bold = True

                # Apply placeholder style if header is empty
                if not header_text or not header_text.strip():
                    run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
                    if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                        run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])

                header_styles = style_kwargs.get("header_styles") or {}
                _apply_cell_style(cell, para, run, header_styles)

        for row_idx, row_data in enumerate(rows):
            table_row = table.add_row()
            cells = table_row.cells
            # Apply row height if specified
            if row_widths and len(row_widths) > row_idx + (1 if headers else 0):
                row_height = row_widths[row_idx + (1 if headers else 0)]
                if row_height:
                    table_row.height = int(row_height * 914400)  # Convert to EMUs
            
            for col_idx, cell_val in enumerate(row_data):
                cell = cells[col_idx]
                para = cell.paragraphs[0]
                # Handle empty cell value
                cell_display = cell_val.strip() if cell_val else DEFAULT_EMPTY_VALUE_TEXT
                run = para.add_run(str(cell_display))
                run.font.bold = False

                # Apply placeholder style if cell is empty
                if not cell_val or not cell_val.strip():
                    run.font.bold = DEFAULT_EMPTY_VALUE_STYLE.get("bold", True)
                    if DEFAULT_EMPTY_VALUE_STYLE.get("font_color"):
                        run.font.color.rgb = RGBColor.from_string(DEFAULT_EMPTY_VALUE_STYLE["font_color"])

                # Apply column styles
                col_styles = (style_kwargs.get("column_styles") or {}).get(col_idx, {})
                _apply_cell_style(cell, para, run, col_styles)

                # Apply row styles
                row_styles = (style_kwargs.get("row_styles") or {}).get(row_idx, {})
                _apply_cell_style(cell, para, run, row_styles)

                # Apply cell-specific styles (highest priority)
                cell_styles = (style_kwargs.get("cell_styles") or {}).get((row_idx, col_idx), {})
                _apply_cell_style(cell, para, run, cell_styles)

        parent.insert(index, table._element)


def _apply_cell_style(cell, para, run, styles):
    """
    Apply styling to a table cell.
    
    This function applies various styling options to a table cell including
    alignment, bold formatting, background color, and font color.
    
    Args:
        cell: The table cell element
        para: The paragraph element within the cell
        run: The text run element within the paragraph
        styles: Dictionary containing styling options:
            - align: Text alignment ("left", "center", "right")
            - bold: Boolean for bold formatting
            - bg_color: Background color as hex string
            - font_color: Font color as hex string
    """
    if styles.get("align"):
        set_paragraph_alignment(para, styles["align"])
    if styles.get("bold"):
        run.font.bold = True
    if styles.get("bg_color"):
        _set_cell_bg_color(cell, styles["bg_color"])
    if styles.get("font_color"):
        run.font.color.rgb = RGBColor.from_string(styles["font_color"])


def _set_cell_bg_color(cell, hex_color):
    """
    Set the background color of a table cell.
    
    Args:
        cell: The table cell element
        hex_color: Background color as hex string (e.g., "FF0000")
    """
    cell_xml = cell._tc
    props = cell_xml.get_or_add_tcPr()
    props.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))
