import os
from docx import Document
from docxblocks.core.inserter import DocxBuilder

def test_table_block(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["A", "B"],
                "rows": [["1", "2"], ["3", "4"]]
            },
            "style": {
                "header_styles": {"bold": True},
                "column_styles": {1: {"align": "center"}},
                "row_styles": {1: {"bg_color": "FFF8DC"}},
                "cell_styles": {(1, 1): {"font_color": "FF0000"}}
            }
        }
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    # Check that table exists and has correct number of rows
    tables = doc2.tables
    assert len(tables) == 1
    assert len(tables[0].rows) == 3  # header + 2 rows


def test_table_block_with_integers(tmp_path):
    """Test that integer values in table cells work correctly without .strip() errors."""
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{main}}")
    doc.save(str(template))

    blocks = [
        {
            "type": "table",
            "content": {
                "headers": ["Name", "Age", "Score"],
                "rows": [
                    ["Alice", 25, 95.5],
                    ["Bob", 30, 87],
                    ["Charlie", 22, 92.8]
                ]
            },
            "style": {
                "header_styles": {"bold": True},
                "column_styles": {1: {"align": "center"}, 2: {"align": "right"}}
            }
        }
    ]
    builder = DocxBuilder(str(template))
    builder.insert("{{main}}", blocks)
    builder.save(str(output))

    assert os.path.exists(output)
    doc2 = Document(str(output))
    tables = doc2.tables
    assert len(tables) == 1
    assert len(tables[0].rows) == 4  # header + 3 rows
    
    # Verify the content is correctly displayed
    table = tables[0]
    # Check headers
    assert table.cell(0, 0).text.strip() == "Name"
    assert table.cell(0, 1).text.strip() == "Age"
    assert table.cell(0, 2).text.strip() == "Score"
    
    # Check data rows
    assert table.cell(1, 0).text.strip() == "Alice"
    assert table.cell(1, 1).text.strip() == "25"
    assert table.cell(1, 2).text.strip() == "95.5"
    
    assert table.cell(2, 0).text.strip() == "Bob"
    assert table.cell(2, 1).text.strip() == "30"
    assert table.cell(2, 2).text.strip() == "87"
    
    assert table.cell(3, 0).text.strip() == "Charlie"
    assert table.cell(3, 1).text.strip() == "22"
    assert table.cell(3, 2).text.strip() == "92.8" 