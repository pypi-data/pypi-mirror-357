#!/usr/bin/env python3
"""
Word Document to Markdown Converter Service

This service handles conversion of Word documents (.docx, .doc) to Markdown format.
"""

import sys
from pathlib import Path
from typing import List, Dict, Optional, Set
import re
import zipfile
import io

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx is not installed. Please install it using:")
    print("pip install python-docx")
    sys.exit(1)

from .base_converter import BaseDocumentConverter


class WordDocumentConverter(BaseDocumentConverter):
    """Converts Word documents to Markdown format with full content preservation."""

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported Word document extensions."""
        return ['.docx', '.doc']
    
    def can_convert(self, file_path: Path) -> bool:
        """Check if this converter can handle Word documents."""
        return file_path.suffix.lower() in self.get_supported_extensions()

    def _clean_formatting_markers(self, text: str) -> str:
        """Clean up consecutive formatting markers in text."""
        # Fix consecutive bold markers: **text****more** -> **text more**
        text = re.sub(r'\*\*([^*]*)\*\*\*\*([^*]*)\*\*', r'**\1\2**', text)

        # Fix consecutive italic markers: *text**more* -> *text more*
        text = re.sub(r'\*([^*]*)\*\*([^*]*)\*', r'*\1\2*', text)

        # Fix mixed markers: **bold***italic* -> ***bold italic***
        text = re.sub(r'\*\*([^*]*)\*\*\*([^*]*)\*', r'***\1\2***', text)

        # Remove empty formatting markers
        text = re.sub(r'\*\*\*\*', '', text)  # Empty bold
        text = re.sub(r'(?<!\*)\*\*(?!\*)', '', text)  # Empty bold (not part of italic)
        text = re.sub(r'(?<!\*)\*(?!\*)', '', text)  # Empty italic

        return text
    
    def _convert_paragraph_to_markdown(self, paragraph: Paragraph, extracted_images: Optional[Dict[str, Path]] = None, processed_images: Optional[Set[str]] = None) -> str:
        """Convert a Word paragraph to Markdown format."""
        text = paragraph.text.strip()

        if processed_images is None:
            processed_images = set()

        # Check for images in this paragraph
        paragraph_images = []
        if extracted_images:
            paragraph_images = self._check_paragraph_for_images(paragraph)

        # If paragraph has no text but has images, skip processing here
        # Images are now handled by position mapping in _convert_document_to_markdown
        if not text and paragraph_images:
            if '__IMAGE_PARAGRAPH__' in paragraph_images:
                # This is an image paragraph, but images are handled by position mapping
                return ""
            else:
                # Handle specific named images (if any)
                markdown_content = ""
                if extracted_images:
                    for img_filename in paragraph_images:
                        if img_filename in extracted_images and img_filename not in processed_images:
                            img_path = extracted_images[img_filename]
                            image_markdown = self._convert_image_to_markdown(img_path)
                            processed_images.add(img_filename)
                            markdown_content += image_markdown
                return markdown_content

        # If paragraph has no text and no images, return empty
        if not text:
            return ""
        
        # Handle different paragraph styles
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # Check if this looks like a heading based on style or content
        is_heading = False
        heading_level = 1
        
        # First check style-based headings
        if 'heading 1' in style_name:
            is_heading = True
            heading_level = 1
        elif 'heading 2' in style_name:
            is_heading = True
            heading_level = 2
        elif 'heading 3' in style_name:
            is_heading = True
            heading_level = 3
        elif 'heading 4' in style_name:
            is_heading = True
            heading_level = 4
        elif 'heading 5' in style_name:
            is_heading = True
            heading_level = 5
        elif 'heading 6' in style_name:
            is_heading = True
            heading_level = 6
        else:
            # Only detect headings if they have clear indicators to preserve original content
            # Check for explicit section numbering patterns first
            section_patterns = [
                r'^(\d+(?:\.\d+)*)\s+(.+)$',  # "4.10 Glossary", "1.2.3 Title" (with space)
            ]

            section_match = None
            for pattern in section_patterns:
                match = re.match(pattern, text)
                if match:
                    section_match = match
                    break

            if section_match:
                # Determine heading level based on number of dots
                section_number = section_match.group(1)
                dot_count = section_number.count('.')
                heading_level = min(dot_count + 1, 6)  # Cap at level 6
                is_heading = True
            else:
                # Be more conservative - only detect as heading if it's clearly formatted as one
                # Check if paragraph is short, all bold, and looks like a title
                if len(text) < 100 and len(text.split()) <= 10:  # Short and concise
                    all_bold = True
                    has_text = False

                    for run in paragraph.runs:
                        if run.text.strip():
                            has_text = True
                            if not run.bold:
                                all_bold = False
                                break

                    # Only treat as heading if it's all bold AND matches common heading patterns
                    if all_bold and has_text:
                        # Check for common document section patterns
                        heading_patterns = [
                            r'^[A-Z][A-Z\s&/]+$',  # All caps like "INTRODUCTION", "BUSINESS REQUIREMENTS"
                            r'^\d+\.\s*[A-Z]',     # Numbered sections like "1. Introduction"
                            r'^[A-Z][a-z]+(\s+[A-Z][a-z]*)*:?$',  # Title case like "Business Requirements"
                        ]

                        is_likely_heading = any(re.match(pattern, text) for pattern in heading_patterns)

                        if is_likely_heading:
                            is_heading = True
                            heading_level = 2  # Default to level 2 for detected headings

        # Convert to markdown heading if identified as heading
        if is_heading:
            heading_prefix = "#" * heading_level
            section_number = self._update_section_counter(heading_level)
            return f"{heading_prefix} {section_number}{text}\n\n"

        # Handle formatting within runs for regular paragraphs
        markdown_text = ""
        for run in paragraph.runs:
            run_text = run.text

            # Apply bold formatting
            if run.bold and run_text.strip():
                run_text = f"**{run_text}**"

            # Apply italic formatting
            if run.italic and run_text.strip():
                run_text = f"*{run_text}*"

            markdown_text += run_text

        # Clean up multiple consecutive formatting markers
        markdown_text = self._clean_formatting_markers(markdown_text)

        # Add any images found in this paragraph inline with the text
        if paragraph_images and extracted_images:
            for img_filename in paragraph_images:
                if img_filename in extracted_images and img_filename not in processed_images:
                    img_path = extracted_images[img_filename]
                    image_markdown = self._convert_image_to_markdown(img_path)
                    processed_images.add(img_filename)
                    markdown_text += "\n" + image_markdown

        return f"{markdown_text}\n\n"
    
    def _extract_cell_content(self, cell) -> str:
        """Extract content from a table cell, handling multiple paragraphs properly."""
        if not cell.paragraphs:
            return ""

        # Join all paragraphs in the cell with a space to keep content in one line
        # This prevents table structure from breaking due to newlines
        cell_parts = []
        for paragraph in cell.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                cell_parts.append(para_text)

        # Join with space and clean up any extra whitespace
        cell_content = " ".join(cell_parts).strip()

        # Replace any remaining newlines with spaces to ensure table integrity
        cell_content = cell_content.replace('\n', ' ').replace('\r', ' ')

        # Clean up multiple consecutive spaces
        cell_content = ' '.join(cell_content.split())

        return cell_content

    def _convert_table_to_markdown(self, table: Table) -> str:
        """Convert a Word table to Azure DevOps Wiki-compatible Markdown table format."""
        if not table.rows:
            return ""

        # Extract table data
        table_data = []
        for row in table.rows:
            row_cells = [self._extract_cell_content(cell) for cell in row.cells]
            table_data.append(row_cells)

        # Use the Azure DevOps compatible table creation method
        return self._create_azure_devops_table(table_data)
    
    def _analyze_document_structure(self, doc) -> None:
        """Analyze document structure to understand heading patterns."""
        self.logger.info("Analyzing Word document structure...")
        
        section_pattern = r'^(\d+(?:\.\d+)*)\s*(.*)$'
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text and len(text) < 200:  # Focus on potential headings
                style_name = paragraph.style.name if paragraph.style else "None"
                
                # Check for section numbering
                match = re.match(section_pattern, text)
                if match:
                    section_num = match.group(1)
                    section_title = match.group(2).strip()
                    self.logger.info(f"Found section: '{section_num}' - '{section_title}' (Style: {style_name})")
                
                # Check for bold formatting
                all_bold = True
                has_text = False
                for run in paragraph.runs:
                    if run.text.strip():
                        has_text = True
                        if not run.bold:
                            all_bold = False
                            break
                
                if all_bold and has_text and 'heading' not in style_name.lower():
                    self.logger.info(f"Potential bold heading: '{text[:50]}...' (Style: {style_name})")

    def _extract_images_from_word_document(self, doc_path: Path) -> Dict[str, Path]:
        """
        Extract embedded images from Word document and save them as temporary files.

        Args:
            doc_path: Path to the Word document

        Returns:
            Dictionary mapping image filenames to their extracted file paths
        """
        extracted_images: Dict[str, Path] = {}

        if not self.extract_images:
            return extracted_images

        try:
            # Word documents are ZIP files, extract images from the media folder
            with zipfile.ZipFile(doc_path, 'r') as docx_zip:
                # List all files in the ZIP
                file_list = docx_zip.namelist()

                # Find image files in word/media/ directory
                image_files = [f for f in file_list if f.startswith('word/media/') and
                             any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.svg'])]

                if image_files:
                    # Create temp directory for images
                    temp_dir = self._create_temp_image_dir()

                    for img_file in image_files:
                        try:
                            # Extract image data
                            img_data = docx_zip.read(img_file)

                            # Create filename from the original path
                            img_filename = Path(img_file).name
                            temp_img_path = temp_dir / img_filename

                            # Write image to temp file
                            with open(temp_img_path, 'wb') as f:
                                f.write(img_data)

                            # Map the original filename to the extracted path
                            extracted_images[img_filename] = temp_img_path
                            self.logger.info(f"Extracted image: {img_filename}")

                        except Exception as e:
                            self.logger.warning(f"Failed to extract image {img_file}: {str(e)}")

                self.logger.info(f"Extracted {len(extracted_images)} images from {doc_path.name}")

        except Exception as e:
            self.logger.error(f"Failed to extract images from {doc_path}: {str(e)}")

        return extracted_images

    def _check_paragraph_for_images(self, paragraph) -> List[str]:
        """
        Check if a paragraph contains embedded images and return their filenames.

        Args:
            paragraph: Word paragraph object

        Returns:
            List of image filenames found in the paragraph
        """
        image_filenames: List[str] = []

        try:
            # Check for drawing elements in the paragraph's XML
            paragraph_xml = paragraph._element.xml

            # Look for image references in the XML
            if 'drawing' in paragraph_xml or 'pic:pic' in paragraph_xml or 'blip' in paragraph_xml:
                # This paragraph likely contains an image
                # For now, we'll use a simpler approach and check if the paragraph has minimal text
                # but contains drawing elements
                text_content = paragraph.text.strip()

                # If paragraph has very little text but contains drawing XML, it's likely an image paragraph
                if len(text_content) < 10 and ('drawing' in paragraph_xml or 'blip' in paragraph_xml):
                    # Mark this as an image paragraph - we'll process all available images
                    # This is a simplified approach that works better than complex XML parsing
                    self.logger.debug(f"Found paragraph with embedded image content")
                    return ['__IMAGE_PARAGRAPH__']  # Special marker for image paragraphs

        except Exception as e:
            self.logger.debug(f"Error checking paragraph for images: {e}")

        return image_filenames

    def _map_image_paragraph_positions(self, doc, extracted_images: Dict[str, Path]) -> Dict[int, List[str]]:
        """
        Create a mapping of paragraph positions to their associated images based on content context.

        Args:
            doc: Word document object
            extracted_images: Dictionary of extracted images

        Returns:
            Dictionary mapping paragraph index to list of image filenames
        """
        image_positions: Dict[int, List[str]] = {}

        if not extracted_images:
            return image_positions

        self.logger.info(f"Mapping {len(extracted_images)} images to paragraph positions using content-based approach")

        # First, extract the content of each image to understand what it contains
        image_contents = {}
        valid_images = {}  # Track images with successful extractions

        for img_filename, img_path in extracted_images.items():
            try:
                # Get a preview of the image content
                image_markdown = self._convert_image_to_markdown(img_path)

                # Check if the extraction was successful
                if image_markdown and image_markdown.strip():
                    # Extract key phrases from the image content
                    content_preview = image_markdown.lower()
                    image_contents[img_filename] = content_preview
                    valid_images[img_filename] = img_path
                    self.logger.info(f"Image {img_filename} contains: {content_preview[:100]}...")
                else:
                    self.logger.info(f"Skipping failed image extraction for mapping: {img_filename}")
            except Exception as e:
                self.logger.warning(f"Failed to preview image {img_filename}: {e}")

        # Update extracted_images to only include valid images for mapping
        extracted_images = valid_images

        # Now scan through paragraphs to find logical placement positions
        for paragraph_index, paragraph in enumerate(doc.paragraphs):
            paragraph_text = paragraph.text.strip().lower()

            # Look for specific content markers that indicate where images should be placed
            if "state transition" in paragraph_text and "review process" in paragraph_text:
                # This is likely where the review process diagram should go
                # Find the image that contains "mid" or "end" and "year" and "review"
                for img_filename, content in image_contents.items():
                    if "mid" in content and "year" in content and "review" in content:
                        image_positions[paragraph_index + 1] = [img_filename]  # Place after the heading
                        self.logger.info(f"Mapped review process image {img_filename} to paragraph {paragraph_index + 1} (after '{paragraph_text[:50]}')")
                        break

            elif "state transition" in paragraph_text and "plan" in paragraph_text:
                # This is likely where the plan submission diagram should go
                # Find the image that contains form/submission workflow
                for img_filename, content in image_contents.items():
                    if img_filename not in [pos[0] for pos in image_positions.values()]:  # Not already mapped
                        if "form" in content or "draft" in content or "submit" in content:
                            image_positions[paragraph_index + 1] = [img_filename]  # Place after the heading
                            self.logger.info(f"Mapped plan submission image {img_filename} to paragraph {paragraph_index + 1} (after '{paragraph_text[:50]}')")
                            break

        # For any remaining unmapped images, place them at logical positions
        unmapped_images = [img for img in extracted_images.keys()
                          if img not in [pos[0] for pos in image_positions.values()]]

        if unmapped_images:
            self.logger.info(f"Placing {len(unmapped_images)} unmapped images at document start")
            # Place remaining images at the beginning, but in a more controlled way
            for i, img_filename in enumerate(unmapped_images):
                # Place at the very beginning, but after any title/header content
                target_position = i * 2  # Space them out
                image_positions[target_position] = [img_filename]
                self.logger.info(f"Mapped remaining image {img_filename} to paragraph {target_position}")

        self.logger.info(f"Final image mapping: {image_positions}")
        return image_positions

    def _convert_document_to_markdown(self, doc_path: Path) -> str:
        """Convert a Word document to Markdown format."""
        try:
            doc = Document(str(doc_path))

            # Reset section counters for new document
            self._reset_section_counters()

            # Analyze document structure for debugging
            self._analyze_document_structure(doc)

            # Extract embedded images first
            extracted_images = self._extract_images_from_word_document(doc_path)

            # Create a mapping of paragraph positions to image paragraphs
            image_paragraph_positions = self._map_image_paragraph_positions(doc, extracted_images)

            # Track processed images to avoid duplicates
            processed_images: Set[str] = set()

            markdown_content = ""
            paragraph_index = 0

            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Find the corresponding paragraph object
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            # Check if this paragraph position has associated images
                            if paragraph_index in image_paragraph_positions:
                                # Process the specific images for this position
                                for img_filename in image_paragraph_positions[paragraph_index]:
                                    if img_filename in extracted_images and img_filename not in processed_images:
                                        img_path = extracted_images[img_filename]
                                        image_markdown = self._convert_image_to_markdown(img_path)
                                        processed_images.add(img_filename)
                                        markdown_content += image_markdown

                            # Convert the paragraph normally
                            paragraph_markdown = self._convert_paragraph_to_markdown(paragraph, extracted_images, processed_images)
                            markdown_content += paragraph_markdown
                            paragraph_index += 1
                            break

                elif element.tag.endswith('tbl'):  # Table
                    # Find the corresponding table object
                    for table in doc.tables:
                        if table._element == element:
                            markdown_content += self._convert_table_to_markdown(table)
                            break

            # Add any remaining unprocessed images at the end
            remaining_images = set(extracted_images.keys()) - processed_images
            if remaining_images:
                markdown_content += "\n\n<!-- Additional embedded images found in document -->\n\n"
                for img_filename in remaining_images:
                    img_path = extracted_images[img_filename]
                    image_markdown = self._convert_image_to_markdown(img_path)
                    markdown_content += image_markdown

            return markdown_content

        except Exception as e:
            self.logger.error(f"Error converting Word document {doc_path}: {str(e)}")
            return f"# Error Converting Document\n\nFailed to convert {doc_path.name}: {str(e)}\n"
