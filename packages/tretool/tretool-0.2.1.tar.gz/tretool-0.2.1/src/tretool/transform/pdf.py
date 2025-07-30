import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Union, List, Optional
from pdfminer.high_level import extract_text
from pdf2image import convert_from_path
from docx import Document
import pandas as pd
from tabula import read_pdf
from PIL import Image
import json


class ConversionError(Exception):
    """自定义转换异常"""
    pass


class PDFConverter(ABC):
    def __init__(self, pdf_path: Union[str, Path]):
        """
        初始化PDF转换器
        
        参数:
            pdf_path: PDF文件路径
        """
        self.pdf_path = Path(pdf_path)
        if not self.pdf_path.exists():
            raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
        if self.pdf_path.suffix.lower() != '.pdf':
            raise ValueError("输入文件必须是PDF格式")

    @abstractmethod
    def convert(self, output_path: Union[str, Path], **kwargs):
        """
        将PDF转换为目标格式
        
        参数:
            output_path: 输出文件路径
            **kwargs: 转换选项
        """
        pass

    @classmethod
    @abstractmethod
    def supported_formats(cls) -> List[str]:
        """返回支持的格式列表"""
        return []

    def _prepare_output_path(self, output_path: Union[str, Path], 
                           default_extension: str) -> Path:
        """
        准备输出路径
        
        参数:
            output_path: 输出路径
            default_extension: 默认文件扩展名
            
        返回:
            处理后的Path对象
        """
        output_path = Path(output_path)
        
        # 如果是目录，自动生成文件名
        if output_path.is_dir():
            output_path = output_path / f"{self.pdf_path.stem}.{default_extension}"
        # 如果没有扩展名，添加默认扩展名
        elif not output_path.suffix:
            output_path = output_path.with_suffix(f".{default_extension}")
            
        # 创建父目录（如果不存在）
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        return output_path


class PDFToDocxConverter(PDFConverter):
    def convert(self, output_path: Union[str, Path], **kwargs):
        """
        将PDF转换为Word文档(.docx)
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - start_page: 开始页(从1开始)
                - end_page: 结束页
                - preserve_formatting: 是否保留格式(True/False)
        """
        try:
            output_path = self._prepare_output_path(output_path, 'docx')
            
            start_page = kwargs.get('start_page', 1)
            end_page = kwargs.get('end_page', None)
            preserve = kwargs.get('preserve_formatting', False)
            
            doc = Document()
            text = extract_text(
                str(self.pdf_path),
                page_numbers=range(start_page-1, end_page) if end_page else None
            )
            
            # 简单的段落处理
            for paragraph in text.split('\n\n'):
                if paragraph.strip():
                    para = doc.add_paragraph()
                    if preserve:
                        # 这里可以添加更复杂的格式保留逻辑
                        runs = paragraph.split('\n')
                        for run in runs:
                            if run.strip():
                                para.add_run(run.strip() + ' ')
                    else:
                        para.add_run(paragraph.strip())
            
            doc.save(output_path)
            return str(output_path)
        except Exception as e:
            raise ConversionError(f"转换为DOCX失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['docx']


class PDFToImageConverter(PDFConverter):
    def convert(self, output_path: Union[str, Path], **kwargs):
        """
        将PDF转换为图像
        
        参数:
            output_path: 输出文件路径或目录
            **kwargs:
                - dpi: 图像DPI(默认200)
                - fmt: 图像格式(png/jpg/tiff)
                - merge: 是否合并所有页为一张长图(True/False)
                - quality: 图像质量(1-100)
        """
        try:
            dpi = kwargs.get('dpi', 200)
            fmt = kwargs.get('fmt', 'png').lower()
            merge = kwargs.get('merge', False)
            quality = kwargs.get('quality', 90)
            
            if fmt not in ['png', 'jpg', 'jpeg', 'tiff']:
                raise ValueError(f"不支持的图像格式: {fmt}")
                
            images = convert_from_path(
                str(self.pdf_path),
                dpi=dpi,
                fmt=fmt if fmt != 'jpg' else 'jpeg'
            )
            
            if merge:
                # 合并所有页为一张长图
                output_path = self._prepare_output_path(output_path, fmt)
                total_height = sum(img.height for img in images)
                max_width = max(img.width for img in images)
                
                merged_image = Image.new('RGB', (max_width, total_height))
                y_offset = 0
                for img in images:
                    merged_image.paste(img, (0, y_offset))
                    y_offset += img.height
                
                merged_image.save(output_path, quality=quality)
                return str(output_path)
            else:
                # 每页保存为单独图像
                output_path = Path(output_path)
                if len(images) == 1:
                    output_path = self._prepare_output_path(output_path, fmt)
                    images[0].save(output_path, quality=quality)
                    return str(output_path)
                else:
                    output_path.mkdir(parents=True, exist_ok=True)
                    output_files = []
                    for i, image in enumerate(images):
                        page_path = output_path / f"page_{i+1}.{fmt}"
                        image.save(page_path, quality=quality)
                        output_files.append(str(page_path))
                    return output_files
        except Exception as e:
            raise ConversionError(f"转换为图像失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['png', 'jpg', 'jpeg', 'tiff']


class PDFToTextConverter(PDFConverter):
    def convert(self, output_path: Union[str, Path], **kwargs):
        """
        将PDF转换为纯文本
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - start_page: 开始页(从1开始)
                - end_page: 结束页
                - encoding: 文本编码(默认utf-8)
        """
        try:
            output_path = self._prepare_output_path(output_path, 'txt')
            
            start_page = kwargs.get('start_page', 1)
            end_page = kwargs.get('end_page', None)
            encoding = kwargs.get('encoding', 'utf-8')
            
            text = extract_text(
                str(self.pdf_path),
                page_numbers=range(start_page-1, end_page) if end_page else None
            )
            
            with open(output_path, 'w', encoding=encoding) as f:
                f.write(text)
            
            return str(output_path)
        except Exception as e:
            raise ConversionError(f"转换为文本失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['txt']


class PDFToCSVConverter(PDFConverter):
    def convert(self, output_path: Union[str, Path], **kwargs):
        """
        提取PDF中的表格为CSV
        
        参数:
            output_path: 输出文件路径或目录
            **kwargs:
                - pages: 要提取的页码('all'或数字或列表)
                - multiple_tables: 如何处理多个表格(separate/merge)
                - encoding: CSV文件编码(默认utf-8)
        """
        try:
            pages = kwargs.get('pages', 'all')
            multiple_tables = kwargs.get('multiple_tables', 'separate')
            encoding = kwargs.get('encoding', 'utf-8')
            
            dfs = read_pdf(str(self.pdf_path), pages=pages, multiple_tables=True)
            
            if not dfs:
                raise ConversionError("未找到表格数据")
                
            if multiple_tables == 'merge':
                # 合并所有表格
                output_path = self._prepare_output_path(output_path, 'csv')
                merged_df = pd.concat(dfs, ignore_index=True)
                merged_df.to_csv(output_path, index=False, encoding=encoding)
                return str(output_path)
            else:
                # 每个表格保存为单独CSV
                output_path = Path(output_path)
                if len(dfs) == 1:
                    output_path = self._prepare_output_path(output_path, 'csv')
                    dfs[0].to_csv(output_path, index=False, encoding=encoding)
                    return str(output_path)
                else:
                    output_path.mkdir(parents=True, exist_ok=True)
                    output_files = []
                    for i, df in enumerate(dfs):
                        table_path = output_path / f"table_{i+1}.csv"
                        df.to_csv(table_path, index=False, encoding=encoding)
                        output_files.append(str(table_path))
                    return output_files
        except Exception as e:
            raise ConversionError(f"提取表格失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['csv']


class PDFToHTMLConverter(PDFConverter):
    def convert(self, output_path: Union[str, Path], **kwargs):
        """
        将PDF转换为HTML
        
        参数:
            output_path: 输出文件路径
            **kwargs:
                - css: 自定义CSS样式
                - images: 是否嵌入图像(True/False)
        """
        try:
            output_path = self._prepare_output_path(output_path, 'html')
            
            # 使用pdfminer提取文本
            text = extract_text(str(self.pdf_path))
            
            # 简单的HTML转换
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{self.pdf_path.stem}</title>
    <style>
        {kwargs.get('css', 'body { font-family: Arial; margin: 20px; }')}
    </style>
</head>
<body>
    <h1>{self.pdf_path.stem}</h1>
    <div id="content">
        {text.replace('\n', '<br>')}
    </div>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return str(output_path)
        except Exception as e:
            raise ConversionError(f"转换为HTML失败: {str(e)}")

    @classmethod
    def supported_formats(cls) -> List[str]:
        return ['html']


class PDFConverterFactory:
    @staticmethod
    def get_converter(target_format: str, pdf_path: str) -> PDFConverter:
        """
        获取指定格式的转换器
        
        参数:
            target_format: 目标格式
            pdf_path: PDF文件路径
            
        返回:
            PDFConverter实例
        """
        format_map = {
            'docx': PDFToDocxConverter,
            'txt': PDFToTextConverter,
            'png': PDFToImageConverter,
            'jpg': PDFToImageConverter,
            'jpeg': PDFToImageConverter,
            'tiff': PDFToImageConverter,
            'csv': PDFToCSVConverter,
            'html': PDFToHTMLConverter,
        }
        
        target_format = target_format.lower()
        if target_format not in format_map:
            raise ValueError(f"不支持的格式: {target_format}")
        
        return format_map[target_format](pdf_path)

    @staticmethod
    def get_supported_formats() -> dict:
        """获取所有支持的格式"""
        return {
            'docx': 'Microsoft Word文档',
            'txt': '纯文本文件',
            'png': 'PNG图像',
            'jpg': 'JPEG图像',
            'jpeg': 'JPEG图像',
            'tiff': 'TIFF图像',
            'csv': 'CSV表格数据',
            'html': 'HTML网页',
        }


# 使用示例
if __name__ == "__main__":
    try:
        # 示例1: PDF转Word
        print("转换PDF到Word...")
        docx_converter = PDFConverterFactory.get_converter('docx', 'example.pdf')
        result = docx_converter.convert('output.docx', preserve_formatting=True)
        print(f"转换成功: {result}")
        
        # 示例2: PDF转图像
        print("\n转换PDF到图像...")
        img_converter = PDFConverterFactory.get_converter('png', 'example.pdf')
        result = img_converter.convert('output_images', dpi=300)
        print(f"转换成功: {result if isinstance(result, str) else len(result)}个文件")
        
        # 示例3: 提取表格数据
        print("\n提取PDF中的表格...")
        csv_converter = PDFConverterFactory.get_converter('csv', 'example.pdf')
        result = csv_converter.convert('output_tables', pages='all')
        print(f"提取成功: {result if isinstance(result, str) else len(result)}个表格")
        
        # 查看所有支持的格式
        print("\n支持的转换格式:")
        for fmt, desc in PDFConverterFactory.get_supported_formats().items():
            print(f"- {fmt}: {desc}")
            
    except Exception as e:
        print(f"发生错误: {str(e)}")