from typing import Any, Dict, Optional
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from AkvoFormPrint.parsers.akvo_flow_parser import AkvoFlowFormParser
from AkvoFormPrint.parsers.akvo_arf_parser import AkvoReactFormParser
from AkvoFormPrint.parsers.default_parser import DefaultParser
from AkvoFormPrint.parsers.base_parser import BaseParser
from AkvoFormPrint.enums import HintText, QuestionType
from AkvoFormPrint.constant import TEXT_ROWS


class DocxRenderer:
    def __init__(
        self,
        orientation: str = "landscape",
        add_section_numbering: bool = False,
        add_question_numbering: bool = True,
        parser_type: Optional[str] = None,
        raw_json: Optional[Dict[str, Any]] = None,
        output_path: Optional[str] = "output.docx",
    ):
        assert orientation in (
            "portrait",
            "landscape",
        ), "Orientation must be 'portrait' or 'landscape'"
        self.orientation = orientation
        self.add_section_numbering = add_section_numbering
        self.add_question_numbering = add_question_numbering
        self.parser_type = parser_type
        self.raw_json = raw_json
        self.output_path = output_path

        self.parser: Optional[BaseParser] = None
        self.form_model = None
        self.doc = Document()

        if parser_type or raw_json:
            self.parser = self._get_parser(parser_type)

        if self.raw_json:
            self.form_model = self._parse_form()

    # ----------------------------- Parsers -----------------------------
    def _get_parser(self, parser_type: Optional[str]) -> BaseParser:
        parsers = {
            "flow": AkvoFlowFormParser,
            "arf": AkvoReactFormParser,
            "default": DefaultParser,
            None: DefaultParser,
        }
        parser_class = parsers.get(parser_type)
        if parser_class:
            return parser_class()
        raise ValueError(f"Unknown parser type: {parser_type}")

    def _parse_form(self):
        if not self.raw_json:
            raise ValueError("No raw_json data provided to parse")
        if not self.parser:
            self.parser = self._get_parser(self.parser_type)
        form_model = self.parser.parse(self.raw_json)
        return self._inject_question_numbers(form_model)

    def _inject_question_numbers(self, form):
        section_index = 0
        counter = 1
        for section in form.sections:
            section.letter = (
                self._number_to_letter(section_index)
                if self.add_section_numbering
                else None
            )
            if self.add_section_numbering:
                section_index += 1
            for question in section.questions:
                question.number = (
                    counter if self.add_question_numbering else None
                )
                if self.add_question_numbering:
                    counter += 1
        return form

    def _number_to_letter(self, n: int) -> str:
        result = ""
        while n >= 0:
            result = chr(n % 26 + ord("A")) + result
            n = n // 26 - 1
        return result

    # ------------------------- Layout Helpers --------------------------
    def _set_landscape(self):
        section = self.doc.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
        for side in [
            "top_margin",
            "bottom_margin",
            "left_margin",
            "right_margin",
        ]:
            setattr(section, side, Inches(0.5))

    def _insert_continuous_section_break(self, paragraph):
        sectPr = OxmlElement("w:sectPr")
        type_elem = OxmlElement("w:type")
        type_elem.set(qn("w:val"), "continuous")
        sectPr.append(type_elem)
        paragraph._p.addnext(sectPr)

    def _set_paragraph_shading_and_underline(
        self,
        paragraph,
        shading: Optional[bool] = True,
        shading_color: Optional[str] = "#F2F2F2",
        underline: Optional[bool] = False,
        underline_color: Optional[str] = "000000",
    ):
        pPr = paragraph._p.get_or_add_pPr()
        if shading:
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), shading_color)
            pPr.append(shd)
        if underline:
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), underline_color)
            pBdr.append(bottom)
            pPr.append(pBdr)

    # ----------------------- Rendering Helpers -------------------------
    def _get_option_symbol_and_values(self, question):
        question_symbol = (
            "( )" if question.type == QuestionType.OPTION else "[ ]"
        )
        options_value = question.answer.options
        if question.answer.allowOther:
            options_value.append("Other:____________________")
        return question_symbol, options_value

    def _add_options_table(self, question):
        option_symbol, options_value = self._get_option_symbol_and_values(
            question
        )

        # calculate left/right option based on options len
        options_len = len(options_value)
        is_odd = options_len % 2 == 1
        col_len = (options_len // 2) + 1 if is_odd else options_len // 2
        col1, col2 = (
            options_value[:col_len],
            options_value[col_len:] if col_len < len(options_value) else [],
        )

        table = self.doc.add_table(rows=1, cols=2)
        table.autofit = True

        for col_idx, col_data in enumerate([col1, col2]):
            cell = table.cell(0, col_idx)
            if col_data:
                first_para = cell.paragraphs[0]
                first_para.text = f"{option_symbol} {col_data[0]}"
                first_para.style.font.size = Pt(10)
                first_para.paragraph_format.space_after = Pt(0)
                first_para.paragraph_format.space_before = Pt(2)
                for opt in col_data[1:]:
                    para = cell.add_paragraph(f"{option_symbol} {opt}")
                    para.style.font.size = Pt(10)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.space_before = Pt(2)

    def _add_single_line_option_table_cell(
        self, question, max_cols: Optional[int] = 3
    ):
        option_symbol, options_value = self._get_option_symbol_and_values(
            question
        )

        options_len = len(options_value)
        cols_len = min(max_cols, options_len)
        rows_len = (options_len + cols_len - 1) // cols_len  # ceil division

        table = self.doc.add_table(rows=rows_len, cols=cols_len)
        table.autofit = True

        # Fill the table left to right, row by row
        for idx, option in enumerate(options_value):
            row = idx // cols_len
            col = idx % cols_len
            cell = table.cell(row, col)
            para = cell.paragraphs[0]
            para.text = f"{option_symbol} {option}"
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(0)
            run = para.runs[0]
            run.font.size = Pt(10)

    def _add_number_boxes_table(self, num_boxes: Optional[int] = 10):
        table = self.doc.add_table(rows=1, cols=num_boxes)

        # Indent the table from the left
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblInd = OxmlElement("w:tblInd")
        tblInd.set(
            qn("w:w"), str(int(Inches(0.15).inches * 1440))
        )  # 0.25 inch indent
        tblInd.set(qn("w:type"), "dxa")
        tblPr.append(tblInd)

        box_width = Inches(0.2)
        for col_idx in range(num_boxes):
            cell = table.cell(0, col_idx)
            # Set fixed width
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(box_width.inches * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            # Set borders
            borders = OxmlElement("w:tcBorders")
            for side in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")
                borders.append(border)
            tcPr.append(borders)

            para = cell.paragraphs[0]
            para.text = ""
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)

    def _add_hint_paragraph(
        self,
        text,
        space_before=0,
        space_after=5,
        add_shading: Optional[bool] = False,
        shading_color: Optional[str] = "#F2F2F2",
    ):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after = Pt(space_after)

        if add_shading:
            self._set_paragraph_shading_and_underline(
                paragraph=para, shading_color=shading_color
            )
        return para

    def _add_horizontal_line(self):
        para = self.doc.add_paragraph()
        p = para._p
        pPr = p.get_or_add_pPr()

        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.5pt line
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")

        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_spacer_paragraph(self):
        spacer_para = self.doc.add_paragraph()
        spacer_para.paragraph_format.space_before = Pt(0)
        spacer_para.paragraph_format.space_after = Pt(0)
        spacer_para.paragraph_format.line_spacing = Pt(1)  # 1pt line height
        # Add a tiny, invisible run (space character)
        run = spacer_para.add_run(" ")
        run.font.size = Pt(1)  # 1pt font, but with a space character

    # -------------------------- Render DOCX ----------------------------
    def render_docx(self):
        if not self.form_model:
            self.form_model = self._parse_form()
        if self.orientation == "landscape":
            self._set_landscape()
        title_para = self.doc.add_paragraph(
            self.form_model.title, style="Title"
        )
        title_para.style.font.size = Pt(14)
        self._insert_continuous_section_break(title_para)

        section = self.doc.sections[-1]
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), "2")
        section._sectPr.append(cols)

        for idx, section_data in enumerate(self.form_model.sections):
            if idx != 0:
                self.doc.add_page_break()
            self._render_section(section_data)

        self.doc.save(self.output_path)

    def _render_section(self, section_data):
        section_title = (
            f"{section_data.letter}. {section_data.title}"
            if section_data.letter
            else section_data.title
        )
        section_para = self.doc.add_paragraph(section_title, style="Heading 1")
        section_para.style.font.size = Pt(12)
        self._set_paragraph_shading_and_underline(
            paragraph=section_para, shading_color="#D9D9D9", underline=True
        )

        for qidx, question in enumerate(section_data.questions):
            if question.type == QuestionType.INSTRUCTION:
                self._render_instruction(qidx, question)
            else:
                self._render_question(qidx, question)

    def _render_question(self, qidx, question):
        required_mark = "*" if question.answer.required else ""
        qtext = (
            f"{question.number}. {question.label} {required_mark}"
            if question.number
            else question.label
        )
        para = self.doc.add_paragraph(qtext)
        para.style.font.size = Pt(10)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(5 if not question.hint else 2)
        self._set_paragraph_shading_and_underline(paragraph=para)

        if question.tooltip:
            self._add_hint_paragraph(text=question.tooltip, add_shading=True)

        if question.hint:
            self._add_hint_paragraph(text=question.hint, add_shading=True)

        if question.dependencies:
            dependency_texts = []
            for d in question.dependencies:
                info = self.form_model.question_id_to_info.get(
                    str(d.depends_on_question_id)
                )
                if info:
                    text = (
                        f'"{d.expected_answer}" selected for '
                        f'question {info[0]}: "{info[1]}"'
                    )
                else:
                    text = (
                        f'"{d.expected_answer}" selected for '
                        f'question "{d.depends_on_question_id}"'
                    )
                dependency_texts.append(text)

            dependency_text = "Answer only if " + " AND ".join(
                dependency_texts
            )
            self._add_hint_paragraph(dependency_text, space_after=2)

        if question.type in [
            QuestionType.OPTION,
            QuestionType.MULTIPLE_OPTION,
        ]:
            if question.answer.optionSingleLine:
                self._add_single_line_option_table_cell(question)
            else:
                self._add_options_table(question)
        elif (
            question.type == QuestionType.NUMBER and question.answer.numberBox
        ):
            self._add_number_boxes_table(question.answer.numberBox)
            self._add_number_hint(question)
        elif question.type == QuestionType.DATE:
            self._render_date_question()
        elif question.type == QuestionType.GEO:
            self._render_geo_question()
        elif question.type == QuestionType.CASCADE:
            self._render_cascade_question(question)
        elif question.type == QuestionType.TEXT:
            textRows = question.answer.textRows or TEXT_ROWS
            for line in range(textRows):
                self._add_horizontal_line()
                if line < textRows - 1:
                    self._add_spacer_paragraph()
        else:
            self._add_spacer_paragraph()
            self._add_horizontal_line()

    def _render_instruction(self, qidx, question):
        para = self.doc.add_paragraph()
        run = para.add_run(question.label)
        run.italic = True
        run.bold = True
        run.font.size = Pt(10)
        para.paragraph_format.space_before = Pt(15)
        para.paragraph_format.space_after = Pt(5)
        self._set_paragraph_shading_and_underline(
            paragraph=para, shading_color="#faebd7"
        )
        if question.tooltip:
            self._add_hint_paragraph(
                text=question.tooltip,
                add_shading=True,
                shading_color="#faebd7",
            )
        self._add_spacer_paragraph()

    def _add_number_hint(self, question):
        min_val, max_val = question.answer.minValue, question.answer.maxValue
        if min_val and max_val:
            text = f"Enter a number between {min_val} and {max_val}"
        elif min_val:
            text = f"Min: {min_val}"
        elif max_val:
            text = f"Max: {max_val}"
        else:
            text = None
        if text:
            self._add_hint_paragraph(text, space_before=4, space_after=0)

    def _render_date_question(self):
        date_para = self.doc.add_paragraph(
            "[    ][    ] / [    ][    ] / [    ][    ][    ][    ]"
        )
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(0)
        self._add_hint_paragraph(
            HintText.DATE.value, space_before=4, space_after=0
        )

    def _render_geo_question(self):
        geo_labels = ["Latitude:", "Longitude:"]
        for lidx, label in enumerate(geo_labels):
            para = self.doc.add_paragraph(label)
            para.style.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(2)
            if lidx == len(geo_labels) - 1:
                para.paragraph_format.space_before = Pt(2)
            else:
                para.paragraph_format.space_before = Pt(0)
            self._add_number_boxes_table(10)

    def _render_cascade_question(self, question):
        for opt in question.answer.options:
            self._add_spacer_paragraph()
            para = self.doc.add_paragraph(f"{opt}: ")
            para.style.font.size = Pt(10)
            self._set_paragraph_shading_and_underline(
                paragraph=para, shading=False, underline=True
            )
