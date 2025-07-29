"""
结构化文档输出模块

该模块提供将结构化文本转换为不同格式文档的功能：
- Markdown (.md)
- 纯文本 (.txt)  
- Word文档 (.docx)
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass
from datetime import datetime
import re

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx未安装，DOCX输出功能不可用")


@dataclass
class DocumentMetadata:
    """文档元数据"""
    title: str = ""
    author: str = ""
    source: str = ""
    created_at: datetime = None
    language: str = "zh"
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now()


@dataclass
class FormattedDocument:
    """格式化文档结果"""
    content: str
    file_path: Path
    format: str
    metadata: DocumentMetadata
    size_bytes: int = 0


class DocumentFormatterError(Exception):
    """文档格式化错误"""
    pass


class DocumentFormatter:
    """文档格式化器"""
    
    SUPPORTED_FORMATS = ['markdown', 'txt', 'docx']
    
    def __init__(self):
        """初始化文档格式化器"""
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def format_document(
        self, 
        content: str, 
        output_path: Union[str, Path], 
        format_type: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> FormattedDocument:
        """
        格式化文档到指定格式
        
        Args:
            content: 要格式化的文本内容
            output_path: 输出文件路径
            format_type: 输出格式 ('markdown', 'txt', 'docx')
            metadata: 文档元数据
            
        Returns:
            FormattedDocument: 格式化结果
        """
        output_path = Path(output_path)
        format_type = format_type.lower()
        
        if format_type not in self.SUPPORTED_FORMATS:
            raise DocumentFormatterError(f"不支持的格式: {format_type}")
        
        if metadata is None:
            metadata = DocumentMetadata()
        
        # 确保输出目录存在
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.logger.info(f"开始格式化文档: {output_path} (格式: {format_type})")
        
        try:
            if format_type == 'markdown':
                return self._format_markdown(content, output_path, metadata)
            elif format_type == 'txt':
                return self._format_txt(content, output_path, metadata)
            elif format_type == 'docx':
                return self._format_docx(content, output_path, metadata)
        except Exception as e:
            raise DocumentFormatterError(f"格式化文档失败: {e}")
    
    def _format_markdown(self, content: str, output_path: Path, metadata: DocumentMetadata) -> FormattedDocument:
        """格式化为Markdown文档"""
        # 构建Markdown内容
        md_content = []
        
        # 添加元数据头部
        if metadata.title:
            md_content.append(f"# {metadata.title}\n")
        
        # 添加元信息
        if any([metadata.author, metadata.source, metadata.created_at]):
            md_content.append("---\n")
            if metadata.author:
                md_content.append(f"**作者**: {metadata.author}\n")
            if metadata.source:
                md_content.append(f"**来源**: {metadata.source}\n")
            if metadata.created_at:
                md_content.append(f"**创建时间**: {metadata.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
            md_content.append("---\n\n")
        
        # 处理内容段落
        formatted_content = self._format_paragraphs(content, 'markdown')
        md_content.append(formatted_content)
        
        final_content = "".join(md_content)
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_content)
        
        size_bytes = output_path.stat().st_size
        self.logger.info(f"Markdown文档已生成: {output_path} ({size_bytes} 字节)")
        
        return FormattedDocument(
            content=final_content,
            file_path=output_path,
            format='markdown',
            metadata=metadata,
            size_bytes=size_bytes
        )
    
    def _format_txt(self, content: str, output_path: Path, metadata: DocumentMetadata) -> FormattedDocument:
        """格式化为纯文本文档"""
        # 构建纯文本内容
        txt_content = []
        
        # 添加标题
        if metadata.title:
            txt_content.append(f"{metadata.title}\n")
            txt_content.append("=" * len(metadata.title) + "\n\n")
        
        # 添加元信息
        if any([metadata.author, metadata.source, metadata.created_at]):
            if metadata.author:
                txt_content.append(f"作者: {metadata.author}\n")
            if metadata.source:
                txt_content.append(f"来源: {metadata.source}\n")
            if metadata.created_at:
                txt_content.append(f"创建时间: {metadata.created_at.strftime('%Y-%m-%d %H:%M:%S')}\n")
            txt_content.append("\n" + "-" * 50 + "\n\n")
        
        # 处理内容段落
        formatted_content = self._format_paragraphs(content, 'txt')
        txt_content.append(formatted_content)
        
        final_content = "".join(txt_content)
        
        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_content)
        
        size_bytes = output_path.stat().st_size
        self.logger.info(f"TXT文档已生成: {output_path} ({size_bytes} 字节)")
        
        return FormattedDocument(
            content=final_content,
            file_path=output_path,
            format='txt',
            metadata=metadata,
            size_bytes=size_bytes
        )
    
    def _format_docx(self, content: str, output_path: Path, metadata: DocumentMetadata) -> FormattedDocument:
        """格式化为DOCX文档"""
        if not DOCX_AVAILABLE:
            raise DocumentFormatterError("python-docx库未安装，无法生成DOCX文档")
        
        # 创建新文档
        document = Document()
        
        # 设置文档属性
        if metadata.title:
            document.core_properties.title = metadata.title
        if metadata.author:
            document.core_properties.author = metadata.author
        document.core_properties.created = metadata.created_at
        
        # 添加标题
        if metadata.title:
            title = document.add_heading(metadata.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元信息表格
        if any([metadata.author, metadata.source, metadata.created_at]):
            table = document.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            if metadata.author:
                row = table.add_row()
                row.cells[0].text = "作者"
                row.cells[1].text = metadata.author
            
            if metadata.source:
                row = table.add_row()
                row.cells[0].text = "来源"
                row.cells[1].text = metadata.source
            
            if metadata.created_at:
                row = table.add_row()
                row.cells[0].text = "创建时间"
                row.cells[1].text = metadata.created_at.strftime('%Y-%m-%d %H:%M:%S')
            
            # 添加分隔符
            document.add_paragraph()
        
        # 处理内容段落
        paragraphs = self._split_into_paragraphs(content)
        
        for paragraph in paragraphs:
            if paragraph.strip():
                p = document.add_paragraph(paragraph.strip())
                # 设置段落格式
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_after = Inches(0.1)
        
        # 保存文档
        document.save(str(output_path))
        
        size_bytes = output_path.stat().st_size
        self.logger.info(f"DOCX文档已生成: {output_path} ({size_bytes} 字节)")
        
        return FormattedDocument(
            content=content,  # DOCX内容不以文本形式返回
            file_path=output_path,
            format='docx',
            metadata=metadata,
            size_bytes=size_bytes
        )
    
    def _format_paragraphs(self, content: str, format_type: str) -> str:
        """格式化段落内容"""
        paragraphs = self._split_into_paragraphs(content)
        
        if format_type == 'markdown':
            # Markdown格式，段落间用双换行分隔
            return "\n\n".join(p.strip() for p in paragraphs if p.strip())
        elif format_type == 'txt':
            # 纯文本格式，段落间用双换行分隔，每段缩进
            formatted_paragraphs = []
            for p in paragraphs:
                if p.strip():
                    # 添加段落缩进
                    formatted_paragraphs.append(f"    {p.strip()}")
            return "\n\n".join(formatted_paragraphs)
        
        return content
    
    def _split_into_paragraphs(self, content: str) -> List[str]:
        """将内容分割为段落"""
        # 按双换行或句号+换行分割段落
        paragraphs = re.split(r'\n\s*\n|(?<=[。！？])\s*\n', content)
        
        # 清理空段落
        return [p.strip() for p in paragraphs if p.strip()]
    
    def batch_format(
        self, 
        content: str, 
        output_dir: Union[str, Path], 
        base_name: str,
        formats: List[str],
        metadata: Optional[DocumentMetadata] = None
    ) -> List[FormattedDocument]:
        """
        批量格式化为多种格式
        
        Args:
            content: 要格式化的内容
            output_dir: 输出目录
            base_name: 基础文件名（不含扩展名）
            formats: 要生成的格式列表
            metadata: 文档元数据
            
        Returns:
            List[FormattedDocument]: 格式化结果列表
        """
        output_dir = Path(output_dir)
        results = []
        
        for format_type in formats:
            if format_type not in self.SUPPORTED_FORMATS:
                self.logger.warning(f"跳过不支持的格式: {format_type}")
                continue
            
            # 确定文件扩展名
            if format_type == 'markdown':
                ext = '.md'
            elif format_type == 'txt':
                ext = '.txt'
            elif format_type == 'docx':
                ext = '.docx'
            
            output_path = output_dir / f"{base_name}{ext}"
            
            try:
                result = self.format_document(content, output_path, format_type, metadata)
                results.append(result)
            except Exception as e:
                self.logger.error(f"生成{format_type}格式失败: {e}")
        
        return results


def create_formatter() -> DocumentFormatter:
    """创建文档格式化器实例"""
    return DocumentFormatter()


def format_document_from_config(
    content: str,
    output_path: Union[str, Path],
    format_type: str,
    config: Optional[Dict] = None
) -> FormattedDocument:
    """
    根据配置格式化文档
    
    Args:
        content: 文档内容
        output_path: 输出路径
        format_type: 格式类型
        config: 配置字典
        
    Returns:
        FormattedDocument: 格式化结果
    """
    formatter = create_formatter()
    
    # 从配置创建元数据
    metadata = DocumentMetadata()
    if config:
        output_config = config.get('output', {})
        metadata.author = output_config.get('author', '')
        metadata.language = output_config.get('language', 'zh')
    
    return formatter.format_document(content, output_path, format_type, metadata) 